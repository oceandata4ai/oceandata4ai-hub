#!/usr/bin/env python3
"""Generate Word copy-review document for expert reviewers."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "copy-review" / "ODA-Demo-Copy-Review.docx"

REVIEW_URL = "https://oceandata4ai.github.io/oceandata4ai-hub/docs/copy-review/review.html"
DEMO_URL = "https://oceandata4ai.github.io/oceandata4ai-hub/"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


SECTIONS = [
    ("A. 全站通用 — 顶部导航", "所有主站页面顶栏", [
        ("G-NAV-01", "Logo", "OceanData4AI", "品牌名"),
        ("G-NAV-02", "主导航", "About · Ask OUG · Blog · Events · Fellows · Join", "导航"),
        ("G-NAV-03", "右侧", "Docs ↗ · GitHub ↗ · Discuss ↗", "CTA"),
        ("G-NAV-04", "主 CTA", "Join Community", "按钮"),
    ]),
    ("A. 全站通用 — Footer", "所有主站页面页脚", [
        ("G-FT-01", "品牌区", "A vendor-neutral community for AI data infrastructure builders.", "正文"),
        ("G-FT-02", "品牌区", "Initiated with support from OceanBase", "免责声明"),
        ("G-FT-07", "底栏", "Built for builders · Not a sales page", "定位语"),
    ]),
    ("B. 首页 Home — Hero & Promo", f"{DEMO_URL}index.html", [
        ("H-02", "Hero H1", "Welcome to the OceanData4AI Community", "标题"),
        ("H-03", "Hero Lead", "Built by builders, for builders. Learn, contribute, connect, and help shape AI-native data infrastructure — from RAG and vectors to agents and hybrid search.", "正文"),
        ("H-04", "Promo 左卡", "Join 800+ builders on Discord", "标题"),
        ("H-05", "Promo 左卡", "Get support, share ideas, and discuss RAG pipelines, vector search, and agent memory with a passionate community.", "正文"),
        ("H-06", "Promo 左卡 CTA", "Join Discord", "按钮"),
        ("H-07", "Promo 右卡", "Stay up to speed with the latest", "标题"),
        ("H-08", "Promo 右卡", "Stay informed on tutorials, deep dives, product releases, and community events — straight from the builders.", "正文"),
        ("H-09", "Promo 右卡 CTA", "Read the Blog", "按钮"),
    ]),
    ("B. 首页 — Fellows / Events / GitHub", DEMO_URL, [
        ("H-10", "Fellows", "Community Fellows Program", "H2"),
        ("H-11", "Fellows", "Get recognized for sharing your expertise on AI data infrastructure with the global builder community.", "正文"),
        ("H-13", "Fellows", "Become a Fellow within the OceanData4AI community. Submit your application showcasing your technical writing skills and experience with RAG, vectors, or agent systems.", "正文"),
        ("H-16", "Fellows CTA", "Apply now", "按钮"),
        ("H-23", "GitHub", "Fork it. Work it. Ship it.", "H2"),
        ("H-24", "GitHub", "Join us on GitHub — contribute code, docs, and examples that help builders ship AI systems faster.", "正文"),
    ]),
    ("B. 首页 — 频道矩阵 Social (6卡)", DEMO_URL, [
        ("H-28", "Discord", "Join 800+ builders for real-time chat, office hours, and peer help.", "描述"),
        ("H-29", "GitHub", "Issues, PRs, and open-source projects from the ecosystem.", "描述"),
        ("H-30", "LinkedIn", "Learn the latest about releases, events, and community news.", "描述"),
        ("H-31", "Medium", "Tutorials, deep dives, and builder stories from the community.", "描述"),
        ("H-32", "Ask OUG", "Searchable Q&A for product questions and deployment help.", "描述"),
        ("H-33", "Events", "Find out when the next meetup, AMA, or workshop will take place.", "描述"),
    ]),
    ("C. About", f"{DEMO_URL}about.html", [
        ("A-01", "Hero", "About OceanData4AI", "标题"),
        ("A-02", "Hero", "Connecting builders who work on the data layer for AI — RAG, vectors, agents, and memory. Vendor-neutral, open, and builder-first.", "正文"),
        ("A-03", "Mission", "OceanData4AI is a community hub for practitioners building AI applications on real data infrastructure…", "正文"),
        ("A-04", "Neutrality", "OceanBase initiated and supports OceanData4AI, but our content covers multiple vendors…", "正文"),
        ("A-07", "Stats", "847 Discord members · 2.4k GitHub stars · 24 Articles published", "数据"),
    ]),
    ("D. Join — 频道入口", f"{DEMO_URL}join.html", [
        ("J-01", "Hero", "Join the community — Pick your path — chat on Discord, contribute on GitHub…", "正文"),
        ("J-02", "Discord", "💬 Discord — fastest way in — Introduce yourself in #introductions… No OceanBase experience required.", "卡片"),
        ("J-03", "GitHub", "⌥ GitHub — contribute code & docs", "卡片"),
        ("J-04", "Events", "📅 Events — meet us live", "卡片"),
        ("J-05", "LinkedIn", "LinkedIn — stay in the loop", "卡片"),
        ("J-06", "Fellow 表单", "Apply as a Fellow — We respond within 5–7 business days.", "表单"),
        ("J-10", "Fellow Hint", "We look for clarity, accuracy, and builder empathy — not marketing copy.", "提示"),
    ]),
    ("E. Blog", f"{DEMO_URL}blog.html", [
        ("B-02", "H1", "Blog", "标题"),
        ("B-03", "Hero", "Tutorials, deep dives, and builder stories on AI data infrastructure — one feed, filtered by topic tags.", "正文"),
        ("B-04", "筛选", "All · News & Updates · Tutorials · Deep Dives · Builder Stories", "标签"),
        ("B-05", "外链", "Read on Medium", "CTA"),
    ]),
    ("F. Events", f"{DEMO_URL}events.html", [
        ("E-01", "H1", "Events", "标题"),
        ("E-02", "Hero", "Office hours, workshops, and AMAs — join us live on Discord.", "正文"),
        ("E-04", "活动", "Tokyo Onchain Night — Invite-only builder meetup in Shibuya…", "卡片"),
        ("E-05", "活动", "Corporate Crypto Strategy Summit — B2B executive session…", "卡片"),
        ("E-07", "Past", "No replays published yet.", "空状态"),
    ]),
    ("G. Fellows", f"{DEMO_URL}fellows.html", [
        ("F-02", "H1", "Fellows Directory", "标题"),
        ("F-03", "Hero", "Builders who publish tutorials, deep dives, and builder stories — selected on merit, not pay-to-play.", "正文"),
        ("F-04", "Featured", "Featured Fellows — Phase 1 showcase", "区标题"),
    ]),
    ("H. Ask OUG — Sidebar & Topics", f"{DEMO_URL}qa/oug-help.html", [
        ("Q-SB-01", "Sidebar", "Discussion", "导航标题"),
        ("Q-SB-02", "Sidebar", "Ask OUG", "导航项"),
        ("Q-SB-03", "Sidebar", "Tags", "标签区标题"),
        ("Q-SB-04", "Sidebar Tag", "seekdb", "标签"),
        ("Q-SB-05", "Sidebar Tag", "deployment", "标签"),
        ("Q-SB-06", "Sidebar Tag", "vector", "标签"),
        ("Q-SB-07", "Sidebar Tag", "hybrid-search", "标签"),
        ("Q-TB-01", "Table header", "Topic · Replies · Views · Activity", "表头"),
        ("Q-03", "H1", "OceanBase User Group Q&A", "标题"),
        ("Q-TP-02", "Topic", "seekdb hybrid search returns empty — knn + BM25 in one SQL", "Demo话题"),
        ("Q-TP-03", "Topic", "How to install OceanBase on Kubernetes (minimal prod checklist)?", "Demo话题"),
    ]),
    ("H. Ask OUG — Signup & Topic detail", f"{DEMO_URL}qa/ask.html", [
        ("Q-08", "Signup", "Sign up for Ask OUG / Sign in to Ask OUG", "标题"),
        ("Q-14", "Verify", "Check your email — Verify email from The OceanData4AI Community", "邮件流"),
        ("Q-16a", "Placeholder", "Summarize your question", "占位"),
        ("Q-TV-05", "Topic page", "Posting and replies open when hosted Q&A goes live…", "说明"),
        ("Q-VF-02", "Verify page", "Email verified — Continue to Ask OUG", "验证成功"),
        ("Q-VF-05", "Verify page", "Verification failed — Back to sign up", "验证失败"),
    ]),
    ("I. M1 Demo 差异", f"{DEMO_URL}m1/", [
        ("M1-01", "Banner", "M1 Milestone Demo · View original hub ↗", "顶栏"),
        ("M1-02", "Hero", "Where builders talk about AI data infrastructure", "H1"),
        ("M1-03", "Gather", "Channels & external matrix — One hub — link out to discuss, publish, and learn.", "区标题"),
    ]),
]


def main():
    doc = Document()
    title = doc.add_heading("OceanData4AI Demo 文案专家评审稿", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run("评审说明\n").bold = True
    for line in [
        f"线上 Demo：{DEMO_URL}",
        f"可视化评审页（左侧文案 + 右侧页面预览）：{REVIEW_URL}",
        "请按 ID 填写文末「专家评审意见」表。",
        "页面位置对照：请打开上方可视化评审页，或运行 scripts/capture-demo-screenshots.sh 生成 PNG。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("文案清单（按页面/频道）", level=1)

    for section_title, page_url, rows in SECTIONS:
        doc.add_heading(section_title, level=2)
        p = doc.add_paragraph()
        run = p.add_run(f"页面预览：{page_url}")
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        run.font.size = Pt(10)
        add_table(
            doc,
            ["ID", "位置", "英文原文", "类型", "评审意见（请填写）"],
            [(*r, "") for r in rows],
        )

    doc.add_page_break()
    doc.add_heading("专家评审意见汇总", level=1)
    add_table(
        doc,
        ["ID", "评审人", "类型", "具体建议", "优先级"],
        [["", "", "措辞/语气/准确性/品牌", "", "P1/P2/P3"]] * 8,
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("完整版 Markdown（含 Blog 文章列表、Fellows bio）：docs/copy-review/ODA-Demo-Copy-Review.md").italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
